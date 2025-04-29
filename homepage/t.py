# from docx import Document
# from docx.oxml.ns import qn
# from docx.oxml import OxmlElement

# def add_strikethrough_to_words(doc_path, words_to_strike):
#     # Load the document
#     doc = Document(doc_path)
    
#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             for word in words_to_strike:
#                 if word in run.text:
#                     # Add strikethrough for the specified word
#                     run.text = run.text.replace(word, word)
#                     run.font.strike = True

#     # Save the modified document
#     doc.save('output.docx')

# # Specify the path to your .docx file
# doc_path = 'example.docx'
# # Specify the words that need to have a center line (strikethrough)
# words_to_strike = ['mananade', 'Hanumanta']

# add_strikethrough_to_words(doc_path, words_to_strike)
from docx import Document
from docx.shared import RGBColor

def add_blue_strikethrough(doc_path, words_to_strike):
    doc = Document(doc_path)
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for word in words_to_strike:
                if word in run.text:
                    # Replace the word with the same text with strikethrough
                    run.text = run.text.replace(word, word)
                    run.font.strike = True  # Apply strikethrough formatting
                    # Note: `python-docx` does not support changing strikethrough color

    doc.save('output.docx')

# Path to your .docx file
doc_path = 'example.docx'
# Words to apply the strikethrough effect
words_to_strike = ['Hanumanta ', 'mananade']

add_blue_strikethrough(doc_path, words_to_strike)
