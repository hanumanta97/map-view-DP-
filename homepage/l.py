# from docx import Document
# from docx.oxml.ns import qn
# from docx.oxml import OxmlElement
# from lxml import etree

# def add_blue_strikethrough(doc_path, words_to_strike):
#     doc = Document(doc_path)
    
#     namespace = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             for word in words_to_strike:
#                 if word in run.text:
#                     # Replace the word with a strikethrough version
#                     run.text = run.text.replace(word, word)
                    
#                     # Access the underlying XML of the run
#                     run_xml = run._element
                    
#                     # Ensure the run properties (rPr) element exists
#                     rPr = run_xml.find('.//w:rPr', namespaces={'w': namespace})
#                     if rPr is None:
#                         rPr = OxmlElement('w:rPr')
#                         run_xml.insert(0, rPr)
                    
#                     # Add the strikethrough element with blue color
#                     strike = OxmlElement('w:strike')
#                     rPr.append(strike)
                    
#                     # Create a blue color element and add it to the run properties
#                     color = OxmlElement('w:color')
#                     color.set(qn('w:val'), '0000FF')  # Blue color in hex
                    
#                     # Add the color to the run's properties
#                     rPr.append(color)

#     # Save the modified document
#     doc.save('output.docx')

# # Path to your .docx file
# doc_path = 'example.docx'
# # Words to apply the strikethrough effect
# words_to_strike = ['Hanumanta', 'mananade']

# add_blue_strikethrough(doc_path, words_to_strike)
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_blue_strikethrough_shape(doc_path, words_to_strike):
    doc = Document(doc_path)

    namespace = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    for paragraph in doc.paragraphs:
        # Create a new run to insert shapes
        new_run = paragraph.add_run()  
        for word in words_to_strike:
            if word in paragraph.text:
                # Find the position of the word
                start_index = paragraph.text.find(word)
                end_index = start_index + len(word)
                
                # Get the XML element of the paragraph
                p_xml = paragraph._element
                
                # Create shape XML for a strikethrough line
                shape_xml = OxmlElement('w:shape')
                shape_xml.set(qn('w:type'), '#_x0000_t75')  # Type of shape (line)
                shape_xml.set(qn('w:strokeColor'), '0000FF')  # Blue color
                shape_xml.set(qn('w:strokeWidth'), '1pt')  # Line width

                # Append the shape to the paragraph
                p_xml.append(shape_xml)

                # Update the text of the run
                new_run.text = word
                break

    # Save the modified document
    doc.save('output.docx')

# Path to your .docx file
doc_path = 'example.docx'
# Words to apply the strikethrough effect
words_to_strike = ['Hanumanta', 'mananade']

add_blue_strikethrough_shape(doc_path, words_to_strike)
