# from docx import Document
# from docx.oxml import OxmlElement
# from lxml import etree
# from docx.oxml.ns import qn

# def add_blue_strikethrough(doc_path, words_to_strike):
#     doc = Document(doc_path)
    
#     # Define XML namespace for Word documents
#     namespace = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             for word in words_to_strike:
#                 if word in run.text:
#                     # Replace the word with a strikethrough version
#                     run.text = run.text.replace(word, word)
                    
#                     # Add strikethrough
#                     run.font.strike = True
                    
#                     # Access the underlying XML of the run
#                     run_xml = run._element
                    
#                     # Create a new element for the strikethrough with blue color
#                     strike = OxmlElement('w:strike')
#                     strike.set(qn('w:val'), 'true')
                    
#                     # Append the strikethrough element to the run
#                     run_xml.insert(0, strike)
                    
#                     # Set the color of the strikethrough (using XML directly)
#                     rPr = run_xml.find('.//w:rPr', namespaces={'w': namespace})
#                     if rPr is None:
#                         rPr = OxmlElement('w:rPr')
#                         run_xml.insert(0, rPr)
                    
#                     # Create a blue color element
#                     color = OxmlElement('w:color')
#                     color.set(qn('w:val'), '0000FF')  # Blue color in hex
                    
#                     # Add the color to the run's properties
#                     rPr.append(color)
    
#     # Save the modified document
#     doc.save('output.docx')

# # Path to your .docx file
# doc_path = 'example.docx'
# # Words to apply the strikethrough effect
# words_to_strike = ['Hanumanta', 'mananade']

# add_blue_strikethrough(doc_path, words_to_strike)
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_blue_strikethrough(doc_path, words_to_strike):
    doc = Document(doc_path)
    
    # Define the namespace
    namespace = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for word in words_to_strike:
                if word in run.text:
                    # Create a new element for the strikethrough
                    run.text = run.text.replace(word, word)
                    run.font.strike = True  # Apply strikethrough
                    
                    # Access the underlying XML of the run
                    run_xml = run._element
                    
                    # Create the color element for strikethrough
                    rPr = run_xml.find('.//w:rPr', namespaces={'w': namespace})
                    if rPr is None:
                        rPr = OxmlElement('w:rPr')
                        run_xml.insert(0, rPr)
                    
                    # Add a strikethrough element
                    strike = OxmlElement('w:strike')
                    rPr.append(strike)
                    
                    # Add color to strikethrough
                    color = OxmlElement('w:color')
                    color.set(qn('w:val'), '0000FF')  # Blue color in hex
                    strike.append(color)

    # Save the modified document
    doc.save('output.docx')

# Path to your .docx file
doc_path = 'example.docx'
# Words to apply the strikethrough effect
words_to_strike = ['Hanumanta', 'mananade']

add_blue_strikethrough(doc_path, words_to_strike)
